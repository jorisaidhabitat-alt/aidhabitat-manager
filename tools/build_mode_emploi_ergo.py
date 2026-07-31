from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "mode-emploi-application-ergo.docx"
APP_ICON = (
    ROOT
    / "aid_habitat_app"
    / "ios"
    / "Runner"
    / "Assets.xcassets"
    / "AppIcon.appiconset"
    / "Icon-App-1024x1024@1x.png"
)

PURPLE = "8F6AA8"
PURPLE_DARK = "684B7A"
PURPLE_LIGHT = "F2ECF5"
PURPLE_PALE = "F8F5FA"
INK = "253043"
MUTED = "667085"
LINE = "D7DCE3"
GREEN = "2F7D58"
GREEN_LIGHT = "EAF5EF"
ORANGE = "B35C00"
ORANGE_LIGHT = "FFF4E5"
RED = "B42318"
RED_LIGHT = "FDECEC"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=160, bottom=100, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color=LINE, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_run_font(run, size=None, color=INK, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_list_numbering(document, ordered):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_list_number(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_paragraph(document, text="", *, bold_prefix=None, color=INK, after=6):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, size=10.5, color=color, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(rest, size=10.5, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=10.5, color=color)
    return paragraph


def add_bullets(document, items, bullet_num_id):
    for item in items:
        paragraph = document.add_paragraph()
        apply_list_number(paragraph, bullet_num_id)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(item)
        set_run_font(run, size=10.5)


def add_steps(document, steps, ordered_num_id):
    ordered_num_id = add_list_numbering(document, ordered=True)
    for step in steps:
        paragraph = document.add_paragraph()
        apply_list_number(paragraph, ordered_num_id)
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(step)
        set_run_font(run, size=10.5)


def add_callout(document, label, text, *, kind="info"):
    palette = {
        "info": (PURPLE_LIGHT, PURPLE_DARK),
        "success": (GREEN_LIGHT, GREEN),
        "warning": (ORANGE_LIGHT, ORANGE),
        "danger": (RED_LIGHT, RED),
    }
    fill, accent = palette[kind]
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=fill, size=0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=180, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, size=10.5, color=accent, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    return paragraph


def add_quick_table(document, rows, widths=(2800, 6560), header=None):
    count = len(rows) + (1 if header else 0)
    table = document.add_table(rows=count, cols=2)
    set_table_geometry(table, list(widths))
    set_table_borders(table)
    row_index = 0
    if header:
        set_repeat_table_header(table.rows[0])
        for index, value in enumerate(header):
            cell = table.rows[0].cells[index]
            set_cell_shading(cell, PURPLE_LIGHT)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(value)
            set_run_font(run, size=10, color=PURPLE_DARK, bold=True)
        row_index = 1
    for label, detail in rows:
        label_cell, detail_cell = table.rows[row_index].cells
        set_cell_shading(label_cell, PURPLE_PALE)
        label_run = label_cell.paragraphs[0].add_run(label)
        set_run_font(label_run, size=10, color=PURPLE_DARK, bold=True)
        detail_run = detail_cell.paragraphs[0].add_run(detail)
        set_run_font(detail_run, size=10, color=INK)
        row_index += 1
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        1: (17, PURPLE_DARK, 16, 8),
        2: (13.5, PURPLE_DARK, 12, 6),
        3: (11.5, INK, 9, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_sections(document):
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("APP’ERGO  ·  MODE D’EMPLOI")
        set_run_font(run, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("Usage interne  ·  Juillet 2026  ·  Page ")
        set_run_font(run, size=9, color=MUTED)
        add_page_field(p)


def add_cover(document):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(58)
    p.paragraph_format.space_after = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if APP_ICON.exists():
        run = p.add_run()
        run.add_picture(str(APP_ICON), width=Inches(1.18))
        drawing = run._r.xpath(".//wp:docPr")
        if drawing:
            drawing[0].set("descr", "Icône de l’application App’Ergo")

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(10)
    run = kicker.add_run("GUIDE TERRAIN")
    set_run_font(run, size=10.5, color=PURPLE, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Mode d’emploi de l’application")
    set_run_font(run, size=27, color=INK, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("Pour les ergothérapeutes Aid’Habitat")
    set_run_font(run, size=15, color=PURPLE_DARK)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Version 1.0 · 30 juillet 2026")
    set_run_font(run, size=10, color=MUTED)


def add_section_break(document):
    document.add_page_break()


def build_document():
    document = Document()
    document.core_properties.title = "Mode d’emploi de l’application App’Ergo"
    document.core_properties.subject = "Guide terrain pour les ergothérapeutes Aid’Habitat"
    document.core_properties.author = "Aid’Habitat"
    document.core_properties.keywords = "App’Ergo, ergothérapie, visite à domicile, mode d’emploi"
    configure_styles(document)
    configure_sections(document)
    ordered_num_id = add_list_numbering(document, ordered=True)
    bullet_num_id = add_list_numbering(document, ordered=False)

    add_cover(document)
    add_section_break(document)

    add_heading(document, "Sommaire", 1)
    add_quick_table(
        document,
        [
            ("1. Démarrer", "Installer, mettre à jour et se connecter."),
            ("2. Se repérer", "Comprendre le menu et accéder à ses dossiers."),
            ("3. Préparer un dossier", "Consulter les informations et les notes avant la visite."),
            ("4. Réaliser le relevé", "Compléter chaque partie de la visite à domicile."),
            ("5. Écrire et dessiner", "Notes, dictée, Apple Pencil, pages et outils."),
            ("6. Documents et photos", "Ajouter, scanner, consulter et annoter."),
            ("7. Ressources", "Bibliothèque, caisses de retraite et portail Anah."),
            ("8. Travailler hors ligne", "Enregistrer localement puis synchroniser."),
            ("9. Générer le rapport", "Contrôler le résumé et produire le PDF."),
            ("10. Dépanner", "Réagir correctement aux messages et signaler un bug."),
        ],
    )
    add_callout(
        document,
        "Lecture rapide",
        "Pour une première prise en main, lire les chapitres 1, 4, 8 et 10. Le reste peut être consulté au besoin.",
        kind="success",
    )

    add_heading(document, "Avant chaque visite", 2)
    add_bullets(
        document,
        [
            "Ouvrir l’application avec une connexion Internet.",
            "Vérifier que le dossier du bénéficiaire est visible.",
            "Ouvrir une fois le dossier et le relevé pour les rendre disponibles hors ligne.",
            "Contrôler la batterie de l’iPad et de l’Apple Pencil.",
            "Attendre la fin d’une synchronisation en cours avant de partir.",
        ],
        bullet_num_id,
    )

    add_section_break(document)
    add_heading(document, "1. Démarrer", 1)
    add_heading(document, "Sur iPad avec TestFlight", 2)
    add_steps(
        document,
        [
            "Ouvrir TestFlight avec le compte Apple invité par Aid’Habitat.",
            "Sélectionner App’Ergo puis toucher Installer.",
            "Lorsqu’une mise à jour est proposée, toucher Mettre à jour avant une journée de visites.",
            "Ouvrir App’Ergo depuis l’écran d’accueil de l’iPad.",
        ],
        ordered_num_id,
    )
    add_callout(
        document,
        "À retenir",
        "Une version TestFlight est mise à jour comme une application classique. Installer les mises à jour proposées permet de recevoir les corrections récentes.",
        kind="info",
    )

    add_heading(document, "Sur Mac ou dans un navigateur", 2)
    add_steps(
        document,
        [
            "Ouvrir https://app.aidhabitat.fr/ dans un navigateur récent.",
            "Vérifier que l’adresse affichée commence bien par https://app.aidhabitat.fr/.",
            "Se connecter avec son profil personnel.",
        ],
        ordered_num_id,
    )

    add_heading(document, "Se connecter", 2)
    add_steps(
        document,
        [
            "Choisir son nom dans la liste des utilisateurs.",
            "Saisir son mot de passe personnel.",
            "Toucher Se connecter.",
        ],
        ordered_num_id,
    )
    add_callout(
        document,
        "Confidentialité",
        "Ne jamais partager son mot de passe ni utiliser le profil d’un collègue. Chaque ergothérapeute voit uniquement les dossiers qui lui sont attribués ; le profil administrateur peut voir l’ensemble des dossiers.",
        kind="warning",
    )

    add_section_break(document)
    add_heading(document, "2. Se repérer dans l’application", 1)
    add_paragraph(
        document,
        "Le menu latéral reste accessible pendant la navigation. Survoler une icône sur Mac ou la toucher sur iPad permet d’identifier sa fonction.",
    )
    add_quick_table(
        document,
        [
            ("Accueil", "Vue synthétique de l’activité et accès aux dossiers récents."),
            ("Dossiers", "Liste des bénéficiaires attribués à l’utilisateur connecté."),
            ("Bibliothèque", "Équipements et exemples d’adaptation disponibles hors ligne."),
            ("Caisses de retraite", "Caisses principales et complémentaires, contacts et modalités."),
            ("Anah", "Accès au portail MaPrimeAdapt’ lorsque la connexion est disponible."),
            ("Réglages", "Profil, photo et gestion des accès pour l’administrateur."),
        ],
        header=("Espace", "Utilité"),
    )
    add_callout(
        document,
        "Navigation mémorisée",
        "Si vous quittez Dossiers pour ouvrir une autre rubrique, un nouveau clic sur Dossiers vous ramène à l’endroit précédemment consulté. Si vous êtes déjà dans Dossiers, un second clic revient à la liste des bénéficiaires.",
        kind="info",
    )

    add_heading(document, "Retrouver un dossier", 2)
    add_steps(
        document,
        [
            "Ouvrir Dossiers.",
            "Attendre la fin du chargement avant de sélectionner un bénéficiaire.",
            "Utiliser la recherche si la liste est longue.",
            "Toucher la carte du bénéficiaire pour ouvrir son dossier.",
        ],
        ordered_num_id,
    )

    add_section_break(document)
    add_heading(document, "3. Préparer le dossier d’un bénéficiaire", 1)
    add_paragraph(
        document,
        "Le dossier rassemble les informations administratives, les notes préparatoires, les documents et l’accès au relevé de visite.",
    )
    add_heading(document, "Informations bénéficiaire", 2)
    add_bullets(
        document,
        [
            "Vérifier l’identité, les coordonnées, l’adresse et la commune.",
            "Contrôler la composition du foyer, les revenus et le statut d’occupation.",
            "Compléter les caisses de retraite principale et complémentaire si elles sont connues.",
            "Ajouter une caisse supplémentaire uniquement lorsqu’un cumul doit être conservé.",
        ],
        bullet_num_id,
    )
    add_callout(
        document,
        "Saisie rapide",
        "Les modifications sont enregistrées localement immédiatement. Pour certains champs, la sauvegarde se déclenche lorsque vous quittez le champ.",
        kind="success",
    )

    add_heading(document, "Notes préparatoires", 2)
    add_paragraph(
        document,
        "La zone de notes est partagée entre une note écrite et une note dessin. La séparation peut être déplacée pour donner davantage de place à l’une ou à l’autre.",
    )
    add_bullets(
        document,
        [
            "Utiliser la note écrite pour les observations utiles avant la visite.",
            "Utiliser la note dessin pour un croquis, un repérage ou une annotation visuelle.",
            "Les notes restent liées au dossier lorsque vous changez de page.",
        ],
        bullet_num_id,
    )

    add_heading(document, "Avant de partir", 2)
    add_steps(
        document,
        [
            "Ouvrir l’espace Documents et vérifier les pièces disponibles.",
            "Ouvrir Visite à domicile : le relevé de visite s’affiche directement.",
            "Parcourir rapidement les onglets pour confirmer qu’ils sont chargés.",
            "Revenir au premier onglet et attendre la fin de la synchronisation.",
        ],
        ordered_num_id,
    )

    add_section_break(document)
    add_heading(document, "4. Réaliser le relevé de visite", 1)
    add_paragraph(
        document,
        "Le relevé est organisé par onglets. Il est possible de revenir sur un onglet à tout moment : les données sont conservées localement au fur et à mesure.",
    )
    add_quick_table(
        document,
        [
            ("Bénéficiaire", "Profil, foyer, santé et informations administratives."),
            ("Contexte de vie", "Habitudes, environnement, autonomie et situation médicale."),
            ("Mesures", "Prise de notes dessin sur les silhouettes assise et debout."),
            ("Accessibilité", "Accès extérieur, niveaux, chauffage, ouvertures et équipements."),
            ("Salle de bain", "Équipements et contraintes par pièce et par niveau."),
            ("WC", "Équipements, hauteur, barres et circulation par pièce."),
            ("Plans", "Plans avant et après travaux, dessin et éléments déplaçables."),
            ("Photos", "Logement, accessibilité, sanitaires et plans."),
            ("Résumé", "Contrôle global des informations et visualisation des notes."),
            ("Préconisations", "Sélection des solutions depuis la bibliothèque et classement."),
        ],
        header=("Onglet", "Contenu"),
        widths=(2300, 7060),
    )

    add_heading(document, "Méthode recommandée", 2)
    add_steps(
        document,
        [
            "Avancer de gauche à droite dans les onglets.",
            "Renseigner les informations certaines et laisser vide ce qui doit être vérifié.",
            "Utiliser les observations pour préciser une situation sans déduire une information médicale.",
            "Ajouter les photos pendant la visite ou ultérieurement au bureau depuis le Mac.",
            "Relire le Résumé avant de terminer.",
        ],
        ordered_num_id,
    )
    add_callout(
        document,
        "Plusieurs occupants",
        "Sélectionner l’occupant concerné avant de compléter une donnée individuelle. Vérifier le prénom affiché dans le bandeau de l’occupant.",
        kind="info",
    )

    add_section_break(document)
    add_heading(document, "5. Écrire, dicter et dessiner", 1)
    add_heading(document, "Note écrite", 2)
    add_bullets(
        document,
        [
            "Toucher l’endroit où le texte doit être ajouté.",
            "Utiliser le clavier, la dictée intégrée ou l’écriture manuscrite Apple Scribble.",
            "La dictée utilise les fonctions Apple de l’appareil ; aucune IA externe n’est nécessaire.",
            "Relire les noms propres, mesures et termes médicaux après une dictée.",
        ],
        bullet_num_id,
    )

    add_heading(document, "Note dessin", 2)
    add_quick_table(
        document,
        [
            ("Crayon", "Écriture et dessin à main levée."),
            ("Surligneur", "Mise en évidence avec une couleur translucide."),
            ("Gomme", "Effacement local ; toucher la gomme pour régler son épaisseur."),
            ("Couleur", "Choix de la couleur actuellement utilisée."),
            ("Annuler / rétablir", "Revenir sur les dernières actions."),
            ("Tout effacer", "Effacer la page après confirmation."),
            ("+ / corbeille", "Ajouter ou supprimer une page après confirmation."),
        ],
        header=("Outil", "Fonction"),
    )
    add_callout(
        document,
        "Apple Pencil",
        "Le double toucher alterne avec l’outil précédemment utilisé. Si vous passez du crayon à la gomme, un nouveau double toucher revient au crayon.",
        kind="info",
    )
    add_callout(
        document,
        "Pages de note",
        "Lorsque vous changez d’onglet puis revenez, l’application conserve la page consultée. Après suppression d’une page, les numéros sont remis dans l’ordre.",
        kind="success",
    )

    add_section_break(document)
    add_heading(document, "6. Plans, photos et préconisations", 1)
    add_heading(document, "Plans", 2)
    add_steps(
        document,
        [
            "Choisir la page de plan.",
            "Indiquer s’il s’agit d’un plan avant ou après travaux.",
            "Dessiner avec les outils de note et utiliser la ligne droite si nécessaire.",
            "Ajouter les éléments proposés, puis les déplacer par glisser-déposer.",
            "Attendre la sauvegarde avant de changer d’onglet.",
        ],
        ordered_num_id,
    )

    add_heading(document, "Photos", 2)
    add_steps(
        document,
        [
            "Toucher le cadre en pointillés de la rubrique concernée.",
            "Choisir une photo existante ou prendre une photo.",
            "Vérifier l’image avant de poursuivre.",
            "Ajouter d’autres photos dans le même cadre si nécessaire.",
        ],
        ordered_num_id,
    )

    add_heading(document, "Préconisations", 2)
    add_steps(
        document,
        [
            "Toucher Ajouter une préconisation.",
            "Choisir un élément de la bibliothèque.",
            "Adapter le titre et la description au projet réel.",
            "Répéter l’opération pour chaque solution retenue.",
            "Maintenir une carte et la déplacer pour définir l’ordre du rapport.",
        ],
        ordered_num_id,
    )
    add_callout(
        document,
        "Important",
        "Une préconisation n’apparaît qu’après la sélection d’un élément de la bibliothèque. Ne pas créer de carte vide.",
        kind="warning",
    )

    add_section_break(document)
    add_heading(document, "7. Gérer les documents", 1)
    add_heading(document, "Ajouter un document", 2)
    add_steps(
        document,
        [
            "Ouvrir Documents dans le dossier du bénéficiaire.",
            "Toucher Déposer un fichier ou prendre une photo.",
            "Choisir Fichier, Photo ou Scanner un document selon le besoin.",
            "Avec le scanner, cadrer le document, conserver la bonne prise puis terminer l’import.",
            "Renommer le document de manière compréhensible.",
        ],
        ordered_num_id,
    )
    add_callout(
        document,
        "Nommage conseillé",
        "Utiliser un nom court et explicite : « Avis d’imposition 2025 », « Plan logement », « Devis salle de bain ».",
        kind="success",
    )

    add_heading(document, "Consulter ou annoter un PDF", 2)
    add_steps(
        document,
        [
            "Toucher la carte du document pour l’ouvrir.",
            "Naviguer jusqu’à la page souhaitée.",
            "Ajouter les annotations nécessaires.",
            "Toucher Enregistrer pour conserver les modifications.",
            "Pour fermer sans conserver, toucher Quitter sans enregistrer.",
        ],
        ordered_num_id,
    )
    add_callout(
        document,
        "Fermeture",
        "Si une modification n’est pas enregistrée, une confirmation doit être affichée. Utiliser la croix pour continuer l’édition.",
        kind="warning",
    )

    add_section_break(document)
    add_heading(document, "8. Utiliser les ressources", 1)
    add_heading(document, "Bibliothèque Wiki et inspirations", 2)
    add_bullets(
        document,
        [
            "Les images sont enregistrées dans l’application et restent visibles hors ligne.",
            "Choisir un filtre pour limiter les résultats.",
            "Une recherche efface le filtre actif et porte sur toute la bibliothèque.",
            "Toucher une carte pour ouvrir sa fiche détaillée.",
            "Toucher en dehors de la fenêtre pour la fermer.",
        ],
        bullet_num_id,
    )

    add_heading(document, "Caisses de retraite", 2)
    add_bullets(
        document,
        [
            "Basculer entre caisses principales et complémentaires par balayage horizontal.",
            "Toucher une caisse pour consulter les profils éligibles, la demande, le délai, le téléphone, le site et le montant.",
            "Toucher un numéro de téléphone pour lancer l’appel.",
            "Dans le dossier bénéficiaire, ajouter plusieurs caisses si nécessaire.",
        ],
        bullet_num_id,
    )

    add_heading(document, "Portail Anah", 2)
    add_paragraph(
        document,
        "Le portail MaPrimeAdapt’ est intégré à l’application native. Il nécessite une connexion Internet et peut être ouvert dans Safari si le site refuse l’affichage intégré.",
    )
    add_callout(
        document,
        "Hors ligne",
        "La rubrique Anah est indisponible sans connexion. Les autres informations déjà chargées restent accessibles.",
        kind="warning",
    )

    add_section_break(document)
    add_heading(document, "9. Travailler hors ligne et synchroniser", 1)
    add_heading(document, "Ce qui se passe automatiquement", 2)
    add_paragraph(
        document,
        "Chaque saisie est d’abord enregistrée sur l’appareil. Dès qu’une connexion est disponible, l’application transmet les opérations en attente vers la base partagée.",
    )
    add_quick_table(
        document,
        [
            ("Aucun bandeau", "La situation est normale ; la synchronisation est terminée ou aucune action n’est en attente."),
            ("Bandeau violet", "Des opérations attendent le réseau ou leur traitement."),
            ("Bandeau rouge", "Une ou plusieurs opérations ont échoué et nécessitent un contrôle."),
        ],
        header=("Affichage", "Signification"),
    )

    add_heading(document, "Avant de passer hors ligne", 2)
    add_steps(
        document,
        [
            "Se connecter avec Internet.",
            "Ouvrir la liste des dossiers et attendre son chargement.",
            "Ouvrir les dossiers prévus pour la journée.",
            "Ouvrir leur relevé de visite et leurs documents utiles.",
            "Vérifier qu’aucun bandeau rouge n’est présent.",
        ],
        ordered_num_id,
    )

    add_heading(document, "Au retour de la connexion", 2)
    add_steps(
        document,
        [
            "Laisser l’application ouverte quelques instants.",
            "Attendre la disparition du bandeau de synchronisation.",
            "Si une génération de rapport était en attente, laisser l’application la reprendre.",
            "Contrôler les dossiers modifiés avant de se déconnecter.",
        ],
        ordered_num_id,
    )
    add_callout(
        document,
        "Ne pas abandonner",
        "Ne jamais toucher Abandonner sur une opération en échec sans consigne. L’abandon retire l’opération de la file mais n’envoie pas la modification au serveur. Ouvrir Détails, prendre une capture sans données sensibles et utiliser Réessayer.",
        kind="danger",
    )

    add_section_break(document)
    add_heading(document, "10. Contrôler et générer le rapport", 1)
    add_heading(document, "Avant la génération", 2)
    add_bullets(
        document,
        [
            "Relire le Résumé et vérifier les informations principales.",
            "Contrôler les plans avant/après travaux.",
            "Vérifier l’ordre et le contenu des préconisations.",
            "S’assurer que les photos utiles sont présentes.",
            "Attendre la fin des sauvegardes en cours.",
        ],
        bullet_num_id,
    )

    add_heading(document, "Générer", 2)
    add_steps(
        document,
        [
            "Toucher Générer en haut du relevé.",
            "Laisser l’application préparer et synchroniser les données.",
            "Attendre l’ouverture ou l’ajout du PDF dans Documents.",
            "Ouvrir le PDF et contrôler les pages importantes.",
        ],
        ordered_num_id,
    )
    add_callout(
        document,
        "Sans connexion",
        "La demande de génération est mise en attente. Elle doit reprendre au retour du réseau. Ne pas multiplier les clics sur Générer.",
        kind="info",
    )

    add_section_break(document)
    add_heading(document, "11. Sécurité et confidentialité", 1)
    add_bullets(
        document,
        [
            "Utiliser uniquement son compte personnel.",
            "Verrouiller l’iPad dès qu’il n’est plus utilisé.",
            "Ne pas transmettre de données bénéficiaires par messagerie personnelle.",
            "Éviter les captures d’écran contenant des informations médicales ou des coordonnées.",
            "Ne pas stocker les rapports dans un espace personnel non autorisé.",
            "Signaler immédiatement la perte ou le vol d’un appareil.",
            "Se déconnecter avant de remettre l’appareil à une autre personne.",
        ],
        bullet_num_id,
    )
    add_callout(
        document,
        "Données sensibles",
        "Les informations de santé, d’autonomie, de revenus et d’identité doivent être consultées uniquement pour l’accompagnement du bénéficiaire.",
        kind="danger",
    )

    add_heading(document, "Checklist de fin de journée", 2)
    add_bullets(
        document,
        [
            "Connexion Internet rétablie.",
            "Aucun bandeau rouge.",
            "Aucune génération de rapport en attente.",
            "Documents et rapports contrôlés.",
            "Application fermée ou session déconnectée si l’appareil est partagé.",
            "iPad remis en charge.",
        ],
        bullet_num_id,
    )

    add_section_break(document)
    add_heading(document, "12. Dépannage rapide", 1)
    add_quick_table(
        document,
        [
            ("Le dossier n’apparaît pas", "Vérifier le profil connecté, la connexion, puis relancer la synchronisation. L’attribution du dossier doit correspondre à l’ergothérapeute."),
            ("Une page reste vide", "Revenir à l’écran précédent, attendre quelques secondes puis rouvrir. Signaler si le problème se répète."),
            ("Bandeau violet permanent", "Vérifier Internet, laisser l’application ouverte, puis se reconnecter si nécessaire."),
            ("Bandeau rouge", "Ouvrir Détails, conserver le message, toucher Réessayer. Ne pas abandonner sans consigne."),
            ("Le rapport reste en attente", "Rétablir Internet, laisser l’application ouverte et se reconnecter si la reprise ne démarre pas."),
            ("Un document ne s’ouvre pas", "Vérifier qu’il a déjà été chargé en ligne. Réessayer avec la connexion."),
            ("Le portail Anah échoue", "Vérifier Internet puis utiliser Ouvrir dans Safari."),
            ("Une note semble revenir en arrière", "Ne plus modifier la zone, attendre la synchronisation et signaler le dossier concerné."),
        ],
        header=("Problème", "Action"),
        widths=(3000, 6360),
    )

    add_heading(document, "Signaler un problème", 2)
    add_steps(
        document,
        [
            "Toucher Signaler dans l’application.",
            "Choisir le type de problème.",
            "Décrire l’action réalisée juste avant le problème.",
            "Envoyer le signalement.",
        ],
        ordered_num_id,
    )
    add_callout(
        document,
        "Message utile",
        "Décrire uniquement le problème observé. Le nom du dossier et l’onglet sont déjà ajoutés automatiquement au signalement. Exemple : « Après l’ajout d’une ligne, un bandeau rouge s’affiche. L’opération n’a pas été abandonnée. »",
        kind="success",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
